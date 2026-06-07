"""
CONSTRUCTOR DE DOCUMENTOS — WORD (.docx) + EXCEL (.xlsx) VINCULADOS
──────────────────────────────────────────────────────────────────
- build_word():  convierte la propuesta (Markdown) en un .docx profesional, con
                 portada, estilos, tablas y una sección que REFERENCIA el Excel y
                 reproduce los totales calculados (mismos números que el Excel).
- build_excel(): genera un .xlsx con cálculos VIVOS (fórmulas =cantidad*costo,
                 =SUM(...)) para Presupuesto, Marco Lógico y Cronograma.

Ambos se construyen desde el mismo FinancialPackage → los montos siempre coinciden.
Si python-docx u openpyxl no están instalados, las funciones devuelven None sin romper.
"""
import re
import math
from datetime import datetime
from pathlib import Path


# ════════════════════════════════════════════════════════════════════════════
#  MEDICIÓN DE EXTENSIÓN (verificación mecánica de formato)
# ════════════════════════════════════════════════════════════════════════════
# Palabras por página aproximadas según interlineado (base 12pt, márgenes estándar)
_WORDS_PER_PAGE = {1.0: 520, 1.15: 470, 1.5: 360, 2.0: 280}


def _words_per_page(format_spec: dict) -> int:
    spacing = float((format_spec or {}).get("line_spacing", 1.5) or 1.5)
    # tomar el valor de tabla más cercano
    base = min(_WORDS_PER_PAGE, key=lambda s: abs(s - spacing))
    wpp = _WORDS_PER_PAGE[base]
    size = float((format_spec or {}).get("font_size", 12) or 12)
    if size:
        wpp = int(wpp * (12.0 / size))
    return max(120, wpp)


def text_stats(text: str, format_spec: dict) -> dict:
    """Cuenta palabras/caracteres y estima páginas; verifica límites del FormatSpec."""
    plain = re.sub(r"[#>*_`|\-]{1,}", " ", text or "")
    words = len(plain.split())
    chars = len(text or "")
    chars_no_spaces = len(re.sub(r"\s+", "", text or ""))
    pages = max(1, math.ceil(words / _words_per_page(format_spec))) if words else 0

    fs = format_spec or {}
    issues = []
    def _v(key):
        v = fs.get(key)
        return v if isinstance(v, (int, float)) and v else None

    if _v("max_words") and words > fs["max_words"]:
        issues.append(f"Excede el máximo de palabras: {words} > {fs['max_words']}")
    if _v("min_words") and words < fs["min_words"]:
        issues.append(f"Por debajo del mínimo de palabras: {words} < {fs['min_words']}")
    if _v("max_chars") and chars > fs["max_chars"]:
        issues.append(f"Excede el máximo de caracteres: {chars} > {fs['max_chars']}")
    if _v("max_pages") and pages > fs["max_pages"]:
        issues.append(f"Excede el máximo de páginas (estimado): {pages} > {fs['max_pages']}")
    if _v("min_pages") and pages < fs["min_pages"]:
        issues.append(f"Por debajo del mínimo de páginas (estimado): {pages} < {fs['min_pages']}")

    return {
        "word_count": words,
        "char_count": chars,
        "char_count_no_spaces": chars_no_spaces,
        "page_estimate": pages,
        "within_limits": not issues,
        "issues": issues,
    }


# ════════════════════════════════════════════════════════════════════════════
#  CÁLCULOS COMPARTIDOS (fuente de verdad para Word y Excel)
# ════════════════════════════════════════════════════════════════════════════
def budget_totals(financial) -> dict:
    items = financial.budget_items if financial else []
    total_solicitado = sum(i.get("fuente_solicitada", 0) for i in items)
    total_contraparte = sum(i.get("contraparte", 0) for i in items)
    total = total_solicitado + total_contraparte
    by_cat: dict = {}
    for i in items:
        cat = i.get("categoria", "Sin categoría")
        by_cat.setdefault(cat, 0.0)
        by_cat[cat] += i.get("fuente_solicitada", 0) + i.get("contraparte", 0)
    return {
        "total_solicitado": total_solicitado,
        "total_contraparte": total_contraparte,
        "total": total,
        "by_category": by_cat,
        "n_items": len(items),
    }


def _money(v, currency="USD") -> str:
    try:
        return f"{currency} {float(v):,.2f}"
    except (ValueError, TypeError):
        return f"{currency} 0.00"


# ════════════════════════════════════════════════════════════════════════════
#  WORD
# ════════════════════════════════════════════════════════════════════════════
_INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")


def _add_runs_with_bold(paragraph, text):
    """Añade texto a un párrafo respetando **negritas** en línea."""
    pos = 0
    for m in _INLINE_BOLD.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.count("|") >= 2


def _is_separator_row(line: str) -> bool:
    s = line.strip().strip("|")
    return bool(s) and all(set(c.strip()) <= set("-: ") for c in s.split("|"))


def _split_table_row(line: str):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def build_word(proposal_md: str, brief, financial, excel_filename: str, out_path: Path):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    except ImportError:
        return None

    fmt = brief.format_spec if brief and isinstance(brief.format_spec, dict) else {}
    font_name = fmt.get("font_name", "Times New Roman")
    font_size = float(fmt.get("font_size", 12) or 12)
    line_spacing = float(fmt.get("line_spacing", 1.5) or 1.5)
    justify = (fmt.get("alignment", "justify") == "justify")

    doc = Document()

    # ── Estilo base (aplica el formato exigido a TODO el documento) ───────
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = line_spacing
    if justify:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(float(fmt.get("margin_top_cm", 2.5) or 2.5))
        section.bottom_margin = Cm(float(fmt.get("margin_bottom_cm", 2.5) or 2.5))
        section.left_margin = Cm(float(fmt.get("margin_left_cm", 3.0) or 3.0))
        section.right_margin = Cm(float(fmt.get("margin_right_cm", 2.5) or 2.5))

    # ── Portada ───────────────────────────────────────────────────────────
    from models.doc_types import get_doc_type
    dt_name = get_doc_type(brief.doc_type_key).name if brief else "Documento"

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(brief.title if brief else "Documento")
    tr.bold = True
    tr.font.size = Pt(min(22, font_size + 8))
    tr.font.color.rgb = RGBColor(0x1F, 0x3B, 0x73)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cite = fmt.get("citation_style")
    sub.add_run(
        f"{dt_name}\n"
        f"{('Estilo de cita: ' + cite + ' · ') if cite else ''}"
        f"{font_name} {font_size:.0f}pt · interlineado {line_spacing}\n"
        f"Ecuador · {datetime.now().strftime('%d/%m/%Y')}"
    ).italic = True

    doc.add_page_break()

    # ── Cuerpo desde Markdown ─────────────────────────────────────────────
    lines = proposal_md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Tablas Markdown
        if _is_table_row(line):
            table_lines = []
            while i < len(lines) and _is_table_row(lines[i]):
                table_lines.append(lines[i])
                i += 1
            _render_md_table(doc, table_lines)
            continue

        # Encabezados
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("#").strip()
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # Reglas horizontales
        if set(stripped) <= set("-*_") and len(stripped) >= 3:
            i += 1
            continue

        # Listas con viñeta
        if stripped.startswith(("- ", "* ", "• ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(p, stripped[2:].strip())
            i += 1
            continue

        # Listas numeradas
        m = re.match(r"^\d+[.)]\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_bold(p, m.group(1))
            i += 1
            continue

        # Párrafo normal
        p = doc.add_paragraph()
        _add_runs_with_bold(p, stripped)
        i += 1

    # ── Sección de vínculo con el Excel ───────────────────────────────────
    if financial and financial.budget_items:
        doc.add_page_break()
        doc.add_heading("Anexo financiero — Vínculo con la hoja de cálculo", level=1)
        link_p = doc.add_paragraph()
        link_p.add_run(
            "Los cálculos detallados del presupuesto, el marco lógico y el cronograma se "
            "presentan en el archivo Excel adjunto: "
        )
        link_p.add_run(excel_filename).bold = True
        link_p.add_run(
            ". Las cifras de la siguiente tabla resumen se generan desde la misma fuente de "
            "datos que el Excel, por lo que ambos documentos son consistentes."
        )

        t = budget_totals(financial)
        cur = financial.currency
        doc.add_heading("Resumen del presupuesto", level=2)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        hdr[0].paragraphs[0].add_run("Concepto").bold = True
        hdr[1].paragraphs[0].add_run("Monto").bold = True
        rows = [
            ("Solicitado al financiador", _money(t["total_solicitado"], cur)),
            ("Contraparte / cofinanciamiento", _money(t["total_contraparte"], cur)),
            ("TOTAL DEL PROYECTO", _money(t["total"], cur)),
        ]
        for concept, amount in rows:
            cells = tbl.add_row().cells
            cells[0].text = concept
            cells[1].text = amount

        if financial.cofinancing_notes:
            doc.add_paragraph(financial.cofinancing_notes).italic = True

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return str(out_path)


def _render_md_table(doc, table_lines):
    rows = [_split_table_row(l) for l in table_lines if not _is_separator_row(l)]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for cidx in range(ncols):
            val = row[cidx] if cidx < len(row) else ""
            val = _INLINE_BOLD.sub(r"\1", val)
            para = cells[cidx].paragraphs[0]
            run = para.add_run(val)
            if ridx == 0:
                run.bold = True


# ════════════════════════════════════════════════════════════════════════════
#  EXCEL
# ════════════════════════════════════════════════════════════════════════════
def build_excel(brief, financial, out_path: Path):
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        return None

    from models.doc_types import get_doc_type
    cur = financial.currency if financial else "USD"
    dt_name = get_doc_type(brief.doc_type_key).name if brief else "Documento"
    # Duración estimada: máximo mes presente en el cronograma, o 12 por defecto
    sched_months = [
        int(m) for s in (financial.schedule_rows if financial else [])
        for m in (s.get("meses", []) or []) if str(m).strip().lstrip("-").isdigit()
    ]
    duration = max(sched_months) if sched_months else 12
    HEAD_FILL = PatternFill("solid", fgColor="1F3B73")
    HEAD_FONT = Font(bold=True, color="FFFFFF")
    TOT_FILL = PatternFill("solid", fgColor="DCE6F1")
    BOLD = Font(bold=True)
    THIN = Side(style="thin", color="BFBFBF")
    BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
    money_fmt = f'"{cur}" #,##0.00'

    wb = Workbook()

    def _style_header(ws, row, ncols):
        for c in range(1, ncols + 1):
            cell = ws.cell(row=row, column=c)
            cell.fill = HEAD_FILL
            cell.font = HEAD_FONT
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = BORDER

    # ── Hoja 1: Resumen ───────────────────────────────────────────────────
    ws = wb.active
    ws.title = "Resumen"
    ws["A1"] = "RESUMEN — CÁLCULOS DEL DOCUMENTO"
    ws["A1"].font = Font(bold=True, size=14, color="1F3B73")
    meta = [
        ("Documento", brief.title if brief else ""),
        ("Tipo", dt_name),
        ("Duración estimada (meses)", duration),
        ("Moneda", cur),
        ("Generado", datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]
    r = 3
    for k, v in meta:
        ws.cell(row=r, column=1, value=k).font = BOLD
        ws.cell(row=r, column=2, value=v)
        r += 1

    # Totales (referencian la hoja Presupuesto con fórmulas)
    r += 1
    ws.cell(row=r, column=1, value="TOTALES DEL PRESUPUESTO").font = Font(bold=True, size=12, color="1F3B73")
    r += 1
    totals_map = [
        ("Solicitado al financiador", "Presupuesto!F{first}:F{last}"),
        ("Contraparte / cofinanciamiento", "Presupuesto!G{first}:G{last}"),
        ("TOTAL DEL PROYECTO", "Presupuesto!H{first}:H{last}"),
    ]

    # ── Hoja 2: Presupuesto ───────────────────────────────────────────────
    wsp = wb.create_sheet("Presupuesto")
    headers = ["Categoría", "Actividad / Concepto", "Unidad", "Cantidad",
               "Costo unitario", "Solicitado", "Contraparte", "Total"]
    for c, h in enumerate(headers, 1):
        wsp.cell(row=1, column=c, value=h)
    _style_header(wsp, 1, len(headers))

    items = financial.budget_items if financial else []
    first_data = 2
    row = first_data
    for it in items:
        wsp.cell(row=row, column=1, value=it.get("categoria", ""))
        wsp.cell(row=row, column=2, value=it.get("actividad", ""))
        wsp.cell(row=row, column=3, value=it.get("unidad", ""))
        wsp.cell(row=row, column=4, value=it.get("cantidad", 0))
        wsp.cell(row=row, column=5, value=it.get("costo_unitario", 0))
        # Solicitado: usa el valor extraído, pero si es 0 y hay cantidad×costo, calcula
        sol = it.get("fuente_solicitada", 0)
        contrap = it.get("contraparte", 0)
        wsp.cell(row=row, column=6, value=sol)
        wsp.cell(row=row, column=7, value=contrap)
        # Total por fila = Solicitado + Contraparte (fórmula viva)
        wsp.cell(row=row, column=8, value=f"=F{row}+G{row}")
        for c in range(1, 9):
            wsp.cell(row=row, column=c).border = BORDER
            if c in (5, 6, 7, 8):
                wsp.cell(row=row, column=c).number_format = money_fmt
        row += 1

    last_data = row - 1 if row > first_data else first_data
    # Fila de totales con SUM
    if items:
        wsp.cell(row=row, column=1, value="TOTAL").font = BOLD
        for c, col in ((6, "F"), (7, "G"), (8, "H")):
            cell = wsp.cell(row=row, column=c, value=f"=SUM({col}{first_data}:{col}{last_data})")
            cell.font = BOLD
            cell.number_format = money_fmt
            cell.fill = TOT_FILL
        for c in range(1, 9):
            cell = wsp.cell(row=row, column=c)
            cell.border = BORDER
            cell.fill = TOT_FILL

    # Anchos
    widths = [22, 40, 12, 10, 16, 16, 16, 16]
    for c, w in enumerate(widths, 1):
        wsp.column_dimensions[get_column_letter(c)].width = w
    if financial and financial.budget_narrative:
        nrow = row + 2
        wsp.cell(row=nrow, column=1, value="Nota presupuestaria:").font = BOLD
        wsp.cell(row=nrow, column=2, value=financial.budget_narrative)

    # Completar fórmulas de totales en Resumen
    for label, ref_tpl in totals_map:
        ws.cell(row=r, column=1, value=label).font = BOLD
        ref = ref_tpl.format(first=first_data, last=last_data)
        cell = ws.cell(row=r, column=2, value=f"=SUM({ref})" if items else 0)
        cell.number_format = money_fmt
        cell.font = BOLD
        r += 1
    ws.column_dimensions["A"].width = 34
    ws.column_dimensions["B"].width = 40

    # ── Hoja 3: Marco Lógico ──────────────────────────────────────────────
    wsl = wb.create_sheet("Marco Lógico")
    lf_headers = ["Nivel", "Resumen narrativo", "Indicador", "Línea base",
                  "Meta", "Fuente de verificación", "Supuestos"]
    for c, h in enumerate(lf_headers, 1):
        wsl.cell(row=1, column=c, value=h)
    _style_header(wsl, 1, len(lf_headers))
    lr = 2
    for rowd in (financial.logframe_rows if financial else []):
        vals = [rowd.get("nivel", ""), rowd.get("resumen_narrativo", ""),
                rowd.get("indicador", ""), rowd.get("linea_base", ""),
                rowd.get("meta", ""), rowd.get("fuente_verificacion", ""),
                rowd.get("supuestos", "")]
        for c, v in enumerate(vals, 1):
            cell = wsl.cell(row=lr, column=c, value=v)
            cell.border = BORDER
            cell.alignment = Alignment(wrap_text=True, vertical="top")
        lr += 1
    for c, w in enumerate([14, 38, 30, 16, 16, 26, 26], 1):
        wsl.column_dimensions[get_column_letter(c)].width = w

    # ── Hoja 4: Cronograma ────────────────────────────────────────────────
    wsc = wb.create_sheet("Cronograma")
    n_months = max(1, min(int(duration), 60))
    headers_c = ["Actividad", "Responsable", "Entregable"] + [f"M{m}" for m in range(1, n_months + 1)]
    for c, h in enumerate(headers_c, 1):
        wsc.cell(row=1, column=c, value=h)
    _style_header(wsc, 1, len(headers_c))
    cr = 2
    MARK = PatternFill("solid", fgColor="70AD47")
    for sched in (financial.schedule_rows if financial else []):
        wsc.cell(row=cr, column=1, value=sched.get("actividad", ""))
        wsc.cell(row=cr, column=2, value=sched.get("responsable", ""))
        wsc.cell(row=cr, column=3, value=sched.get("entregable", ""))
        months = sched.get("meses", []) or []
        for m in months:
            try:
                mi = int(m)
            except (ValueError, TypeError):
                continue
            if 1 <= mi <= n_months:
                cell = wsc.cell(row=cr, column=3 + mi, value="●")
                cell.fill = MARK
                cell.alignment = Alignment(horizontal="center")
        for c in range(1, len(headers_c) + 1):
            wsc.cell(row=cr, column=c).border = BORDER
        cr += 1
    wsc.column_dimensions["A"].width = 38
    wsc.column_dimensions["B"].width = 20
    wsc.column_dimensions["C"].width = 26
    for m in range(1, n_months + 1):
        wsc.column_dimensions[get_column_letter(3 + m)].width = 5

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))
    return str(out_path)
