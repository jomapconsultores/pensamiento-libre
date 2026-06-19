# -*- coding: utf-8 -*-
"""Genera MAPA_DEL_SISTEMA.docx: arquitectura completa + costos de agente_map."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

AZUL = RGBColor(0x1F, 0x3A, 0x5F)
GRIS = RGBColor(0x55, 0x55, 0x55)
VERDE = RGBColor(0x1B, 0x5E, 0x20)
ROJO = RGBColor(0x9A, 0x2A, 0x2A)

doc = Document()

# ── Estilos base ──
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def h1(txt):
    p = doc.add_heading(txt, level=1)
    for r in p.runs:
        r.font.color.rgb = AZUL
    return p

def h2(txt):
    p = doc.add_heading(txt, level=2)
    for r in p.runs:
        r.font.color.rgb = AZUL
    return p

def para(txt, italic=False, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p

def bullet(txt, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(txt)
    return p

def tabla(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

# ════════════════════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("MAPA DEL SISTEMA\nagente_map")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = AZUL

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Sistema multiagente multiusuario para producir entregables de alto nivel\n"
                "(propuestas de financiamiento no reembolsable, artículos científicos, TDR, tesis, etc.)")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = GRIS

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Cómo está por funcionar y cuánto puede costar\n"
                 "Documento técnico-ejecutivo · 19 de junio de 2026")
r.font.size = Pt(11)
r.font.color.rgb = GRIS

doc.add_paragraph()

# ── Resumen ejecutivo ──
h1("1. Resumen ejecutivo")
para("agente_map es una aplicación web (API HTTP + interfaz) que toma una idea, un tema, un "
     "archivo o una URL, y genera de forma autónoma un entregable profesional terminado en "
     "Word y Excel, sometiéndolo a una verificación de calidad muy exigente (regla 90/90).")
para("El sistema combina VARIOS modelos de inteligencia artificial que trabajan como un equipo:")
bullet("encuentra/analiza la oportunidad de financiamiento e investiga en la web,", "Un agente \"analista\" (Claude) ")
bullet("redactan y mejoran el documento por turnos rotando entre ellos,", "Tres agentes \"constructores\" (Mistral, Codestral, DeepSeek) ")
bullet("revisan entre sí hasta coincidir en que todo está cumplido,", "Los mismos tres constructores ")
bullet("da el veredicto final y exige ≥90 puntos en cada criterio y ≥90 global.", "Un agente \"revisor\" (Claude) ")
para("Está preparado para subir a la nube (Render) con base de datos (Supabase), login con "
     "huella/Face ID y aprobación de usuarios por un administrador.")

p = doc.add_paragraph()
r = p.add_run("Conclusión de costos: ")
r.bold = True; r.font.color.rgb = VERDE
p.add_run("la infraestructura puede arrancar en ~US$0/mes (planes gratuitos) y subir a "
          "~US$7–32/mes al profesionalizar. El costo real variable es el de las APIs de IA: "
          "del orden de US$1 a US$3 por cada propuesta completa generada (estimado). Ver sección 7.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 2. ARQUITECTURA
# ════════════════════════════════════════════════════════════
h1("2. Arquitectura general")
para("El sistema tiene cuatro capas:")

tabla(
    ["Capa", "Qué hace", "Tecnología"],
    [
        ["Interfaz / API", "Recibe peticiones del usuario, login, descarga de documentos",
         "FastAPI + Uvicorn; página web en api/static/index.html (login con clave + WebAuthn huella/Face ID)"],
        ["Orquestador", "Coordina a los agentes en fases y aplica la regla 90/90",
         "core/pipeline.py (versión headless) y main.py (CLI interactiva)"],
        ["Agentes de IA", "Clasifican, investigan, redactan, revisan y estructuran finanzas",
         "agents/*.py — Claude (Anthropic) + Mistral/Codestral/DeepSeek"],
        ["Persistencia / salida", "Guarda sesiones y usuarios; arma Word/Excel a pedido",
         "Supabase (PostgreSQL) + python-docx + openpyxl"],
    ],
    widths=[3.0, 6.5, 7.0],
)
para("La búsqueda web la realiza DuckDuckGo (librería ddgs), sin costo y sin clave de API.",
     italic=True, color=GRIS, size=10)

h2("2.1 Componentes principales (mapa de carpetas)")
tabla(
    ["Módulo", "Función"],
    [
        ["api/main.py", "API HTTP: endpoints de auth, propuestas, descarga Word/Excel, gestión de usuarios"],
        ["api/auth.py", "Hash de contraseñas, tokens de sesión firmados, WebAuthn (huella/Face ID)"],
        ["core/pipeline.py", "Pipeline completo sin interacción (lo que corre en el servidor)"],
        ["agents/analyst.py", "Agente 1 — Analista de viabilidad + buscador de oportunidades (Claude)"],
        ["agents/classifier.py", "Agente 0 — Detecta el tipo de documento y arma el \"brief\""],
        ["agents/writer.py", "Agente 2 — Redactor (constructor rota: Mistral→Codestral→DeepSeek)"],
        ["agents/peer.py", "Consenso de los 3 constructores antes de pasar a Claude"],
        ["agents/reviewer.py", "Agente 3 — Revisor de calidad (Claude), aplica regla 90/90"],
        ["agents/financial.py", "Agente 4 — Estructura presupuesto, marco lógico y cronograma"],
        ["agents/llm.py", "Capa que habla con todos los proveedores (con reintentos y respaldo)"],
        ["tools/search.py", "Motor de búsqueda profunda (DuckDuckGo + descarga de páginas)"],
        ["tools/document_builder.py", "Genera el .docx y el .xlsx finales"],
        ["db/ + utils/supabase_client.py", "Lectura/escritura en Supabase (sesiones, usuarios)"],
    ],
    widths=[5.0, 11.5],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 3. FLUJO
# ════════════════════════════════════════════════════════════
h1("3. Cómo funciona, paso a paso")
para("Cuando un usuario envía una solicitud (POST /propuestas), el trabajo corre en segundo plano "
     "y su estado pasa por: pendiente → en proceso → aprobado / fallido. El flujo es:")

h2("Fase 0 — Clasificación (Claude)")
bullet("Detecta automáticamente el tipo de entregable (propuesta, artículo científico, TDR, tesis, "
       "documento legal/técnico, peer-review o genérico) o respeta el que el usuario eligió.")

h2("Fase 1 — Investigación web + análisis (DeepSeek)  →  revisa Mistral")
bullet("Una IA NO-Claude (DeepSeek) hace la búsqueda en web: propone consultas, ejecuta DuckDuckGo "
       "(gratis), descarga las bases reales e interpreta la evidencia para producir el análisis.")
bullet("Si es propuesta: calcula viabilidad y probabilidad de ganar y exige fuentes con URL real "
       "(no inventa fechas ni montos). Si no es viable, se detiene.")
bullet("GATE: otra IA (Mistral) audita la investigación y le pone puntaje 0-100. Si es <90 → vuelve al inicio.")

h2("Fase 2 — Redacción (DeepSeek)  →  revisa Mistral")
bullet("DeepSeek redacta el documento completo (secciones, formato y lineamientos) a partir del brief. "
       "Es la fase con más tokens de salida, por eso se usa la IA más económica.")
bullet("GATE: otra IA (Mistral) audita la redacción 0-100. Si es <90 → vuelve al inicio (reinvestiga).")

h2("Fase 3 — Estructuración financiera (Codestral)")
bullet("Si el documento requiere presupuesto, Codestral extrae presupuesto, marco lógico y cronograma "
       "en datos estructurados que alimentan el Excel con fórmulas vivas.")

h2("Fase 4 — Veredicto final 90/90 (Claude)")
bullet("Claude da el veredicto final con la regla 90/90: aprueba solo si cada criterio ≥90 y el global "
       "≥90, y el largo está dentro de los límites.")
bullet("Si no aprueba → vuelve al inicio. El ciclo completo se reintenta hasta 5 veces; si nunca llega "
       "a 90, se entrega la mejor versión lograda marcada como inconclusa.")

h2("Salida")
bullet("Se guarda la sesión en Supabase y, a pedido, se generan el Word y el Excel "
       "(no se almacenan en disco: se regeneran on-demand desde la base).")

p = doc.add_paragraph()
r = p.add_run("Regla 90/90: ")
r.bold = True; r.font.color.rgb = ROJO
p.add_run("es el control de calidad central. El revisor reprueba si cualquier criterio individual "
          "queda por debajo de 90, aunque el promedio sea alto. Esto eleva la calidad pero también "
          "puede aumentar el número de ciclos (y, por tanto, el costo en APIs).")

# ── Diagrama de flujo visual ──
if os.path.exists("FLUJO_VISUAL.png"):
    doc.add_page_break()
    h2("3.1 Flujo visual del sistema")
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture("FLUJO_VISUAL.png", width=Inches(6.2))

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 4. AGENTES Y MODELOS
# ════════════════════════════════════════════════════════════
h1("4. Quién hace qué: agentes y modelos de IA")
para("Principio: cada fase la PRODUCE una IA y la AUDITA otra distinta. Asignación por defecto "
     "(configurable por variables de entorno ROLE_*):")
tabla(
    ["Fase", "Tarea", "IA que la hace", "La revisa (gate ≥90)"],
    [
        ["0 · Clasificación", "Detecta el tipo de documento", "Claude (sonnet-4-6)", "—"],
        ["1 · Investigación", "Búsqueda web + análisis de viabilidad", "DeepSeek", "Mistral"],
        ["2 · Redacción", "Escribe el documento completo", "DeepSeek", "Mistral"],
        ["3 · Financiero", "Presupuesto, marco lógico, cronograma", "Codestral", "(veredicto final)"],
        ["4 · Veredicto final", "Regla 90/90 (cada criterio ≥90 y global ≥90)", "Claude (sonnet-4-6)", "—"],
    ],
    widths=[3.4, 5.2, 4.0, 3.9],
)
para("Si cualquier gate da <90, el pipeline VUELVE AL INICIO (reinvestiga), hasta 5 intentos; si no "
     "alcanza el 90, entrega la mejor versión como inconclusa. La búsqueda web la ejecuta DuckDuckGo "
     "(gratis) en código; la IA decide las consultas e interpreta la evidencia. Si un proveedor está "
     "caído, el sistema reintenta con otro y, en último caso, usa Claude como respaldo.",
     italic=True, color=GRIS, size=10)

# ════════════════════════════════════════════════════════════
# 5. ESTADO ACTUAL
# ════════════════════════════════════════════════════════════
h1("5. Estado actual: qué está listo y qué falta")

h2("Listo / implementado")
for t in [
    "Pipeline multiagente completo (clasificar → analizar → construir → consenso → revisar → financiero).",
    "Regla de calidad 90/90 calculada en código (no depende solo del modelo).",
    "API HTTP con FastAPI: crear propuesta, consultar estado, reintentar, descargar Word/Excel, ver historial de revisiones.",
    "Multiusuario: registro, login, aprobación por administrador, propiedad de cada entregable por usuario.",
    "Login con clave + huella/Face ID (WebAuthn) e interfaz web profesional.",
    "Rotación y respaldo entre proveedores de IA para resistir caídas.",
    "Búsqueda web profunda sin costo (DuckDuckGo) con lectura real de bases de convocatorias.",
    "Generación de Word y Excel con formato y cálculos vivos.",
    "Configuración de despliegue en Render (render.yaml) y guía DEPLOY.md.",
]:
    bullet(t)

h2("Pendiente / a verificar antes de operar")
for t in [
    "Cargar las claves de API en el servidor: ANTHROPIC_API_KEY y, opcionalmente, MISTRAL/CODESTRAL/DEEPSEEK.",
    "Crear el proyecto en Supabase y aplicar las migraciones SQL (incluida 004 de usuarios/propiedad).",
    "Confirmar las claves de Supabase (URL, publishable, secret, service-role) en el entorno de Render.",
    "Definir el primer usuario administrador (el primer registro queda como admin aprobado).",
    "Plan Free de Render: el servicio \"se duerme\" tras 15 min sin uso y un pipeline largo podría cortarse "
    "si Render reinicia el proceso. Para producción seria conviene el plan Starter (US$7/mes) o un worker dedicado.",
    "Sin las claves de Mistral/Codestral/DeepSeek, los constructores caen en Claude como respaldo: "
    "funciona, pero sube el costo (todo lo haría Claude).",
]:
    bullet(t)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 6. PRECIOS DE REFERENCIA
# ════════════════════════════════════════════════════════════
h1("6. Precios de referencia de las IAs")
para("Los modelos cobran por \"tokens\" (fragmentos de texto; ~750 palabras ≈ 1.000 tokens). "
     "Se cobra por separado lo que entra (input) y lo que sale (output).")

h2("6.1 Anthropic — Claude (precio exacto, por 1 millón de tokens)")
tabla(
    ["Modelo", "Entrada (US$/1M)", "Salida (US$/1M)", "Uso en el sistema"],
    [
        ["claude-sonnet-4-6 (por defecto)", "$3.00", "$15.00", "Clasificador, Analista, Revisor"],
        ["claude-opus-4-8 (opcional, +calidad)", "$5.00", "$25.00", "Si se sube ANTHROPIC_MODEL"],
        ["claude-haiku-4-5 (más barato)", "$1.00", "$5.00", "Alternativa económica"],
    ],
    widths=[5.5, 3.7, 3.7, 4.0],
)
para("El sistema usa claude-sonnet-4-6 por defecto. Subir a Opus mejora la calidad pero casi duplica el costo.",
     italic=True, color=GRIS, size=10)

h2("6.2 Constructores no-Claude (aproximado — verificar en cada proveedor)")
tabla(
    ["Modelo", "Entrada (US$/1M aprox.)", "Salida (US$/1M aprox.)"],
    [
        ["DeepSeek (deepseek-chat)", "~$0.30–0.60", "~$1.10–1.70"],
        ["Mistral Large", "~$2.00", "~$6.00"],
        ["Codestral", "~$0.30", "~$0.90"],
    ],
    widths=[6.0, 5.5, 5.5],
)
para("⚠ Estos precios cambian con frecuencia y dependen del plan. Confírmalos en console.mistral.ai y "
     "platform.deepseek.com antes de presupuestar. En conjunto, los constructores son notablemente más "
     "baratos que Claude.", italic=True, color=ROJO, size=10)

h2("6.3 Búsqueda web")
bullet("DuckDuckGo (ddgs): GRATIS, sin clave de API. Es la que usa el sistema hoy.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 7. COSTO POR PROPUESTA
# ════════════════════════════════════════════════════════════
h1("7. Costos: cuánto puedes gastar")

h2("7.1 Costo variable por propuesta completa (estimado)")
para("Una propuesta completa puede implicar: 1 análisis con varias rondas de búsqueda, 1–3 ciclos de "
     "redacción, consenso de 3 constructores por ciclo y 1–3 revisiones de Claude. Estimación por "
     "propuesta generada:")
tabla(
    ["Componente", "Quién", "Costo estimado"],
    [
        ["Investigación web + análisis", "DeepSeek", "$0.03 – $0.12"],
        ["Gates de fase (investigación + redacción)", "Mistral", "$0.05 – $0.20"],
        ["Redacción (1+ intentos)", "DeepSeek", "$0.04 – $0.20"],
        ["Estructuración financiera", "Codestral", "$0.03 – $0.10"],
        ["Veredicto final 90/90 (1+ intentos)", "Claude Sonnet", "$0.20 – $0.70"],
        ["TOTAL aproximado por propuesta", "—", "≈ $0.40 – $1.50"],
    ],
    widths=[6.0, 5.0, 4.5],
)
p = doc.add_paragraph()
r = p.add_run("Rango práctico: US$1 a US$3 por propuesta. ")
r.bold = True
p.add_run("Sube cuando: el documento es largo, se llega a los 5 ciclos de revisión, o se usa Opus en "
          "lugar de Sonnet. Baja cuando: el documento es corto y se aprueba en el primer o segundo ciclo.")

para("Si NO configuras las claves de Mistral/Codestral/DeepSeek, todo lo hace Claude (respaldo). En ese "
     "caso el costo por propuesta puede acercarse a US$3–US$6, porque la redacción larga (16.000 tokens "
     "por ciclo) pasaría a tarifas de Claude.", italic=True, color=ROJO, size=10)

h2("7.2 Proyección de volumen (con Claude Sonnet + constructores configurados)")
tabla(
    ["Propuestas/mes", "Costo IA bajo (~$1)", "Costo IA alto (~$3)"],
    [
        ["10", "~$10", "~$30"],
        ["50", "~$50", "~$150"],
        ["100", "~$100", "~$300"],
        ["300", "~$300", "~$900"],
    ],
    widths=[5.0, 5.5, 5.5],
)

h2("7.3 Costos fijos de infraestructura (mensual)")
tabla(
    ["Servicio", "Plan gratis", "Plan recomendado producción"],
    [
        ["Render (servidor web)", "US$0 (Free, se duerme)", "US$7/mes (Starter, siempre activo)"],
        ["Supabase (base de datos)", "US$0 (Free, 500 MB)", "US$25/mes (Pro) si crece el uso"],
        ["Búsqueda (DuckDuckGo)", "US$0", "US$0"],
        ["Dominio propio (opcional)", "—", "~US$1/mes (~US$12/año)"],
        ["TOTAL fijo", "US$0/mes", "≈ US$7 – US$33/mes"],
    ],
    widths=[5.0, 5.5, 6.0],
)

h2("7.4 Costo total estimado de operación")
tabla(
    ["Escenario", "Fijo/mes", "IA (variable)", "Total mensual aprox."],
    [
        ["Pruebas (gratis, ~10 prop.)", "$0", "~$10–30", "≈ $10 – $30"],
        ["Operación pequeña (50 prop.)", "$7", "~$50–150", "≈ $57 – $157"],
        ["Operación media (100 prop.)", "$32", "~$100–300", "≈ $132 – $332"],
        ["Operación alta (300 prop.)", "$32", "~$300–900", "≈ $332 – $932"],
    ],
    widths=[5.2, 3.2, 4.0, 4.1],
)
para("Cifras estimativas para planificación. El gasto real depende del largo de los documentos, del "
     "número de ciclos 90/90 y de los precios vigentes de cada proveedor.", italic=True, color=GRIS, size=10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 8. RIESGOS Y RECOMENDACIONES
# ════════════════════════════════════════════════════════════
h1("8. Riesgos de costo y recomendaciones")

h2("Qué dispara el costo")
bullet("Llegar a los 5 ciclos de revisión por documentos que no superan el 90 en algún criterio.", "Ciclos 90/90: ")
bullet("Cambiar a Opus multiplica ~1,7× el costo de las partes hechas por Claude.", "Modelo Opus: ")
bullet("Sin claves de constructores, la redacción larga la asume Claude (más caro).", "Falta de constructores: ")
bullet("Documentos muy largos = más tokens de entrada al revisar (hasta 60.000 caracteres por revisión).", "Documentos largos: ")

h2("Cómo controlar el gasto")
bullet("Mantén ANTHROPIC_MODEL en claude-sonnet-4-6 (no Opus) salvo casos premium.")
bullet("Configura las claves de Mistral/Codestral/DeepSeek para que la redacción NO la pague Claude.")
bullet("Considera bajar MAX_REVIEW_CYCLES (hoy 5) a 3 si los costos suben de más.")
bullet("Empieza en planes gratuitos (Render Free + Supabase Free) para validar; sube a Starter cuando haya uso real.")
bullet("Fija un tope mensual de gasto en la consola de Anthropic para evitar sorpresas.")
bullet("Lleva control por usuario (ya hay propiedad de entregables) para saber quién consume más.")

h2("Riesgos técnicos a vigilar")
bullet("Plan Free de Render: si el proceso se reinicia durante un pipeline largo, ese trabajo puede quedar en \"en proceso\".")
bullet("Dependes de servicios externos (Anthropic, Mistral, DeepSeek, DuckDuckGo): caídas o cambios de precio te afectan.")
bullet("Verifica límites de tasa (rate limits) si procesas muchas propuestas en paralelo.")

# ── Cierre ──
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("En síntesis")
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = AZUL
para("El sistema está funcionalmente completo y listo para desplegar; lo que falta es configuración "
     "(claves de API y base de datos) y decidir el plan de hosting. El costo de arranque puede ser US$0 "
     "y el costo real está dominado por las APIs de IA: del orden de US$1–US$3 por propuesta con la "
     "configuración recomendada. Con constructores económicos activados y Sonnet (no Opus), el sistema "
     "es claramente sostenible incluso con cientos de propuestas al mes.")

doc.save("MAPA_DEL_SISTEMA.docx")
print("OK -> MAPA_DEL_SISTEMA.docx")
